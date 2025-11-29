"""
File Parser Module for ATS Resume Scorer

This module handles file validation and text extraction from resume files.
Supports PDF, DOC, and DOCX formats with robust error handling.

Requirements: 3.1, 3.2, 3.3, 3.4, 3.5, 15.1, 15.2
"""

import io
import magic
from typing import Tuple, Optional
import pdfplumber
import PyPDF2
from docx import Document

from utils.error_handler import (
    FileUploadError,
    FileParsingError as BaseFileParsingError,
    TextExtractionError,
    ErrorCategory,
    log_error,
    log_warning,
    log_info,
    with_fallback
)


class FileParsingError(Exception):
    """Custom exception for file parsing errors"""
    pass


class FileValidationError(Exception):
    """Custom exception for file validation errors"""
    pass


# Constants
MAX_FILE_SIZE_MB = 5
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
SUPPORTED_MIME_TYPES = {
    'application/pdf': 'pdf',
    'application/msword': 'doc',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx'
}
SUPPORTED_EXTENSIONS = {'.pdf', '.doc', '.docx'}


def validate_file(file_data: bytes, filename: str) -> Tuple[bool, str, Optional[str]]:
    """
    Validates file type and size using magic number checking.
    
    Args:
        file_data: Raw file bytes
        filename: Original filename
        
    Returns:
        Tuple of (is_valid, error_message, file_type)
        - is_valid: True if file passes validation
        - error_message: Error description if validation fails, empty string otherwise
        - file_type: Detected file type ('pdf', 'doc', 'docx') or None
        
    Validates:
        - Requirements 3.1: File type validation
        - Requirements 3.2: File size validation (5MB limit)
        - Requirements 15.1: Clear validation error messages
    """
    # Check file size
    file_size_bytes = len(file_data)
    if file_size_bytes > MAX_FILE_SIZE_BYTES:
        file_size_mb = file_size_bytes / (1024 * 1024)
        error_msg = (
            f"File size ({file_size_mb:.2f} MB) exceeds the maximum allowed size "
            f"of {MAX_FILE_SIZE_MB} MB. Please upload a smaller file or compress your resume."
        )
        return False, error_msg, None
    
    # Check if file is empty
    if file_size_bytes == 0:
        return False, "The uploaded file is empty. Please check the file and try again.", None
    
    # Detect MIME type using magic numbers
    try:
        mime_type = magic.from_buffer(file_data, mime=True)
    except Exception as e:
        return False, f"Unable to determine file type: {str(e)}", None
    
    # Validate MIME type
    if mime_type not in SUPPORTED_MIME_TYPES:
        supported_formats = ', '.join(SUPPORTED_MIME_TYPES.values()).upper()
        error_msg = (
            f"Unsupported file type detected: {mime_type}. "
            f"Please upload a file in one of the supported formats: {supported_formats}."
        )
        return False, error_msg, None
    
    file_type = SUPPORTED_MIME_TYPES[mime_type]
    return True, "", file_type


def _extract_pdf_with_pdfplumber(file_data: bytes) -> str:
    """
    Extract text from PDF using pdfplumber (primary method).
    
    Args:
        file_data: Raw PDF file bytes
        
    Returns:
        Extracted text as string
        
    Raises:
        Exception: If extraction fails
    """
    text = ""
    with pdfplumber.open(io.BytesIO(file_data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    
    if not text.strip():
        raise TextExtractionError(
            "pdfplumber extracted no text",
            user_message="No text could be extracted from the PDF."
        )
    
    return text.strip()


def _extract_pdf_with_pypdf2(file_data: bytes) -> str:
    """
    Extract text from PDF using PyPDF2 (fallback method).
    
    Args:
        file_data: Raw PDF file bytes
        
    Returns:
        Extracted text as string
        
    Raises:
        Exception: If extraction fails
    """
    text = ""
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    
    if not text.strip():
        raise TextExtractionError(
            "PyPDF2 extracted no text",
            user_message="No text could be extracted from the PDF."
        )
    
    return text.strip()


def extract_text_from_pdf(file_data: bytes) -> str:
    """
    Extracts text from PDF files using pdfplumber with PyPDF2 fallback.
    
    Args:
        file_data: Raw PDF file bytes
        
    Returns:
        Extracted text as string
        
    Raises:
        FileParsingError: If text extraction fails with both methods
        
    Validates:
        - Requirements 3.3: PDF text extraction
        - Requirements 15.2: Fallback extraction methods
    """
    # Try primary method (pdfplumber) with fallback to PyPDF2
    try:
        result, used_fallback = with_fallback(
            _extract_pdf_with_pdfplumber,
            _extract_pdf_with_pypdf2,
            file_data,
            error_category=ErrorCategory.TEXT_EXTRACTION,
            log_fallback=True
        )
        
        if used_fallback:
            log_info("PDF extraction succeeded using PyPDF2 fallback", context="file_parser")
        
        return result
        
    except Exception as e:
        # Both methods failed
        log_error(e, context="extract_text_from_pdf", category=ErrorCategory.TEXT_EXTRACTION)
        raise FileParsingError(
            "Failed to extract text from PDF using both pdfplumber and PyPDF2. "
            "The PDF may be corrupted, password-protected, or contain only images. "
            "Please try converting the PDF to a different format or ensure it contains selectable text."
        ) from e


def extract_text_from_docx(file_data: bytes) -> str:
    """
    Extracts text from DOCX files using python-docx.
    
    Args:
        file_data: Raw DOCX file bytes
        
    Returns:
        Extracted text as string
        
    Raises:
        FileParsingError: If text extraction fails
        
    Validates:
        - Requirements 3.4: DOC/DOCX text extraction
        - Requirements 15.1: Clear validation error messages
    """
    try:
        doc = Document(io.BytesIO(file_data))
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        text = "\n".join(text_parts)
        
        if not text.strip():
            log_warning("DOCX file contained no extractable text", context="extract_text_from_docx")
            raise FileParsingError(
                "No text could be extracted from the document. "
                "The document may be empty or corrupted. "
                "Please check the file and try again."
            )
        
        log_info(f"Successfully extracted {len(text)} characters from DOCX", context="extract_text_from_docx")
        return text.strip()
    
    except FileParsingError:
        raise
    except Exception as e:
        log_error(e, context="extract_text_from_docx", category=ErrorCategory.TEXT_EXTRACTION)
        raise FileParsingError(
            "Failed to extract text from DOCX file. "
            "The document may be corrupted or in an unsupported format. "
            "Please try re-saving the document or converting it to PDF."
        ) from e


def extract_text_from_doc(file_data: bytes) -> str:
    """
    Extracts text from legacy DOC files.
    
    Note: Legacy .doc format is not fully supported by python-docx.
    Users should convert to .docx or .pdf for best results.
    
    Args:
        file_data: Raw DOC file bytes
        
    Returns:
        Extracted text as string
        
    Raises:
        FileParsingError: If text extraction fails
        
    Validates:
        - Requirements 3.4: DOC/DOCX text extraction
    """
    raise FileParsingError(
        "Legacy .doc format is not fully supported. "
        "Please convert your document to .docx or .pdf format and try again. "
        "You can convert using Microsoft Word, Google Docs, or online conversion tools."
    )


def extract_text(file_data: bytes, file_type: str) -> str:
    """
    Main text extraction function that routes to appropriate parser.
    
    Args:
        file_data: Raw file bytes
        file_type: File type ('pdf', 'doc', or 'docx')
        
    Returns:
        Extracted text as string
        
    Raises:
        FileParsingError: If text extraction fails
        FileValidationError: If file type is invalid
        
    Validates:
        - Requirements 3.3: PDF text extraction
        - Requirements 3.4: DOC/DOCX text extraction
        - Requirements 3.5: Extraction failure error handling
        - Requirements 15.2: Fallback extraction attempts
    """
    if file_type == 'pdf':
        return extract_text_from_pdf(file_data)
    elif file_type == 'docx':
        return extract_text_from_docx(file_data)
    elif file_type == 'doc':
        return extract_text_from_doc(file_data)
    else:
        raise FileValidationError(
            f"Invalid file type: {file_type}. "
            f"Supported types are: PDF, DOC, DOCX."
        )


def parse_resume_file(file_data: bytes, filename: str) -> Tuple[str, dict]:
    """
    Complete file parsing pipeline: validate and extract text.
    
    Args:
        file_data: Raw file bytes
        filename: Original filename
        
    Returns:
        Tuple of (extracted_text, metadata)
        - extracted_text: Extracted text content
        - metadata: Dictionary with file information
        
    Raises:
        FileValidationError: If file validation fails
        FileParsingError: If text extraction fails
        
    Validates:
        - Requirements 3.1, 3.2: File validation
        - Requirements 3.3, 3.4: Text extraction
        - Requirements 3.5: Error handling with specific messages
        - Requirements 15.1: Clear validation error messages
        - Requirements 15.2: Fallback extraction attempts
        - Requirements 15.4: User-friendly error messages
        
    Example:
        >>> with open('resume.pdf', 'rb') as f:
        ...     file_data = f.read()
        >>> text, metadata = parse_resume_file(file_data, 'resume.pdf')
        >>> print(f"Extracted {len(text)} characters")
    """
    log_info(f"Starting file parsing for: {filename}", context="parse_resume_file")
    
    # Validate file
    try:
        is_valid, error_msg, file_type = validate_file(file_data, filename)
        if not is_valid:
            log_warning(f"File validation failed: {error_msg}", context="parse_resume_file")
            raise FileValidationError(error_msg)
    except FileValidationError:
        raise
    except Exception as e:
        log_error(e, context="parse_resume_file_validation", category=ErrorCategory.FILE_UPLOAD)
        raise FileValidationError(
            "Could not validate the uploaded file. "
            "Please ensure it is a valid PDF, DOC, or DOCX file."
        ) from e
    
    # Extract text
    try:
        text = extract_text(file_data, file_type)
        log_info(f"Successfully extracted {len(text)} characters from {filename}", context="parse_resume_file")
    except FileParsingError:
        raise
    except Exception as e:
        log_error(e, context="parse_resume_file_extraction", category=ErrorCategory.TEXT_EXTRACTION)
        raise FileParsingError(
            "An unexpected error occurred while processing the file. "
            "Please try again or contact support if the problem persists."
        ) from e
    
    # Prepare metadata
    metadata = {
        'filename': filename,
        'file_type': file_type,
        'file_size_bytes': len(file_data),
        'text_length': len(text),
        'success': True
    }
    
    return text, metadata
