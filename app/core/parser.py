"""
Resume File Parser Module
Handles file validation and text extraction from PDF, DOCX, and DOC files.

📌 TEACHING NOTE — What does this file do?
    This is the ENTRY POINT for any uploaded resume file.
    Before the app can analyze a resume, it must:
    1. VALIDATE the file (right size? right type? not empty?)
    2. EXTRACT text from it (read the bytes and turn them into a string)

    Different file formats (PDF, DOCX, DOC) require completely different
    libraries to read — this file handles all of them.

📌 TEACHING NOTE — The two custom exceptions at the top:
    class FileParsingError(Exception): pass
    class FileValidationError(Exception): pass

    These are CUSTOM EXCEPTION TYPES — they inherit from Python's built-in
    Exception but have meaningful names.

    Why custom exceptions vs just raising ValueError or RuntimeError?
    - They make error handling specific: you can catch FileParsingError
      without accidentally catching unrelated errors
    - They communicate intent: the name alone tells you what went wrong
    - Different parts of the app can catch different exception types

📌 TEACHING NOTE — Library overview:
    magic        → detects file type from raw bytes (more reliable than file extension)
    pdfplumber   → primary PDF text extractor (handles complex layouts well)
    PyPDF2       → fallback PDF extractor (simpler, used if pdfplumber fails)
    python-docx  → extracts text from .docx Word documents
"""

import io
import magic
from typing import Tuple, Optional
import pdfplumber
import PyPDF2
from docx import Document

# Custom error classes from the app's utils module
from app.utils.errors import (
    FileUploadError, FileParsingError as BaseFileParsingError,
    TextExtractionError, ErrorCategory,
    log_error, log_warning, log_info, with_fallback
)


# ── Custom Exception Classes ──────────────────────────────────────────────────

class FileParsingError(Exception):
    """
    Raised when a file can be opened but text cannot be extracted from it.

    📌 TEACHING NOTE — Exception class body = 'pass':
        class FileParsingError(Exception): pass

        The 'pass' keyword means "do nothing additional."
        We're creating a new exception TYPE by inheriting from Exception,
        but we don't need any new methods — just the new name is enough.

        This is one of the few valid uses of 'pass' in Python.
        The class has all of Exception's behaviour (message, traceback etc.)
        plus its own identity that allows specific catching:
            try: ...
            except FileParsingError: ...  # catches ONLY this type
    """
    pass


class FileValidationError(Exception):
    """
    Raised when a file fails validation before we even attempt to parse it.
    (wrong size, wrong type, empty file)
    """
    pass


# ── Module-Level Constants ────────────────────────────────────────────────────

# 📌 TEACHING NOTE — Constants at module level:
#   Define magic numbers as named constants at the top of the file.
#   MAX_FILE_SIZE_MB = 5  is readable and easy to change.
#   If you see '5242880' deep in the code, you have to mentally calculate
#   what it means. Named constants = self-documenting code.

MAX_FILE_SIZE_MB    = 5
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024  # 5 × 1024 × 1024 = 5,242,880 bytes

# Maps MIME types (from the 'magic' library) to simple file type strings
SUPPORTED_MIME_TYPES = {
    'application/pdf':   'pdf',
    'application/msword': 'doc',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx'
}

# Valid file extensions for quick display (not used for validation — MIME is)
SUPPORTED_EXTENSIONS = {'.pdf', '.doc', '.docx'}


# ── File Validation ───────────────────────────────────────────────────────────

def validate_file(file_data: bytes, filename: str) -> Tuple[bool, str, Optional[str]]:
    """
    Validate an uploaded file before attempting to extract text from it.

    📌 TEACHING NOTE — Return type: Tuple[bool, str, Optional[str]]
        This function returns THREE values packed into a tuple:
        (is_valid, error_message, file_type)

        is_valid:      True if validation passed, False if not
        error_message: Empty string '' if valid, user-facing error if not
        file_type:     'pdf', 'docx', or 'doc' if valid, None if not

        Python allows returning multiple values as a tuple:
            return True, '', 'pdf'
        The caller unpacks them:
            is_valid, error_msg, file_type = validate_file(data, name)

        This is more convenient than returning a dict for a small fixed set
        of return values.

    📌 TEACHING NOTE — MIME type detection with python-magic:
        File extensions can be faked (.exe renamed to .pdf).
        The 'magic' library reads the first few bytes of the file
        (called "magic bytes" or "file signature") to determine the
        actual file type — regardless of the extension.

        Examples:
        PDF files start with: %PDF (0x25 0x50 0x44 0x46)
        DOCX files start with: PK  (0x50 0x4B — it's actually a ZIP!)

        magic.from_buffer(file_data, mime=True) returns a string like
        'application/pdf' or 'application/zip'.

        This is more secure and reliable than trusting the filename.

    📌 TEACHING NOTE — Validation order (fail fast on cheapest checks):
        1. File size (O(1) — just read len())
        2. Empty file check (O(1))
        3. MIME type detection (slightly more expensive)
        We check simple/cheap things first so we fail fast when possible.

    Args:
        file_data: Raw bytes of the uploaded file
        filename: Original filename (used for error messages)

    Returns:
        Tuple of (is_valid, error_message, file_type_or_None)
    """
    # ── Check 1: File size ────────────────────────────────────────────────
    file_size_bytes = len(file_data)
    if file_size_bytes > MAX_FILE_SIZE_BYTES:
        file_size_mb = file_size_bytes / (1024 * 1024)
        error_msg = (
            f'File size ({file_size_mb:.2f} MB) exceeds the maximum allowed size of '
            f'{MAX_FILE_SIZE_MB} MB. Please upload a smaller file or compress your resume.'
        )
        return False, error_msg, None  # Tuple return: fail immediately

    # ── Check 2: Empty file ───────────────────────────────────────────────
    if file_size_bytes == 0:
        return False, 'The uploaded file is empty. Please check the file and try again.', None

    # ── Check 3: MIME type (actual file type, not just the extension) ─────
    try:
        mime_type = magic.from_buffer(file_data, mime=True)
    except Exception as e:
        return False, f'Unable to determine file type: {str(e)}', None

    if mime_type not in SUPPORTED_MIME_TYPES:
        supported_formats = ', '.join(SUPPORTED_MIME_TYPES.values()).upper()
        error_msg = (
            f'Unsupported file type detected: {mime_type}. '
            f'Please upload a file in one of the supported formats: {supported_formats}.'
        )
        return False, error_msg, None

    # All checks passed — return the file type string
    file_type = SUPPORTED_MIME_TYPES[mime_type]   # e.g., 'application/pdf' → 'pdf'
    return True, '', file_type   # Empty error string = no error


# ── PDF Text Extraction ───────────────────────────────────────────────────────

def _extract_pdf_with_pdfplumber(file_data: bytes) -> str:
    """
    Extract text from a PDF using pdfplumber (primary method).

    📌 TEACHING NOTE — io.BytesIO() — in-memory file:
        Libraries like pdfplumber expect a file object, not raw bytes.
        io.BytesIO(file_data) wraps raw bytes in a file-like object
        that pdfplumber can read — WITHOUT writing to disk.

        This is important for web apps:
        - Writing to disk requires permissions and cleanup
        - In-memory is faster and simpler
        - BytesIO behaves exactly like an open file but lives in RAM

        Pattern used throughout this file:
            with pdfplumber.open(io.BytesIO(file_data)) as pdf:

    📌 TEACHING NOTE — 'with' statement (context manager):
        with pdfplumber.open(...) as pdf:
        This ensures the PDF file is properly closed when the block ends,
        even if an exception is raised. It's equivalent to:
            pdf = pdfplumber.open(...)
            try:
                ... work ...
            finally:
                pdf.close()
        Always use 'with' for files, database connections, etc.

    📌 TEACHING NOTE — Building text page by page:
        text += page_text + '\n'
        We accumulate text from each page into one string.
        The '\n' ensures page text doesn't run together.
        if page_text: guard skips empty pages (some PDFs have blank pages).

    Args:
        file_data: Raw PDF bytes

    Returns:
        Extracted text as a single string

    Raises:
        TextExtractionError: If no text was found (e.g., scanned image PDF)
    """
    text = ''
    with pdfplumber.open(io.BytesIO(file_data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:   # Skip pages with no text (e.g., image pages)
                text += page_text + '\n'

    if not text.strip():
        raise TextExtractionError(
            'pdfplumber extracted no text',
            user_message='No text could be extracted from the PDF.'
        )
    return text.strip()


def _extract_pdf_with_pypdf2(file_data: bytes) -> str:
    """
    Extract text from a PDF using PyPDF2 (fallback method).

    📌 TEACHING NOTE — Why two PDF libraries?
        pdfplumber (primary): Better at complex layouts, tables, columns.
        PyPDF2 (fallback): Simpler, faster, handles some files pdfplumber can't.

        Neither library works on ALL PDFs perfectly.
        A PDF that pdfplumber fails on might work with PyPDF2 and vice versa.
        By trying both, we maximize the chance of successfully extracting text.

        This is the FALLBACK PATTERN — try the best option first,
        fall back to alternatives before giving up entirely.

    📌 TEACHING NOTE — PyPDF2 vs pdfplumber API difference:
        pdfplumber: pdf.pages (iterator of page objects)
        PyPDF2:     pdf_reader.pages (same concept, different class names)

        Both iterate pages and call extract_text(), but the underlying
        PDF parsing algorithms differ significantly.

    Args:
        file_data: Raw PDF bytes

    Returns:
        Extracted text as string

    Raises:
        TextExtractionError: If no text extracted
    """
    text = ''
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + '\n'

    if not text.strip():
        raise TextExtractionError(
            'PyPDF2 extracted no text',
            user_message='No text could be extracted from the PDF.'
        )
    return text.strip()


def extract_text_from_pdf(file_data: bytes) -> str:
    """
    Public PDF extraction function — tries pdfplumber then falls back to PyPDF2.

    📌 TEACHING NOTE — with_fallback() utility:
        with_fallback(primary_fn, fallback_fn, *args) is a custom utility
        that runs primary_fn(*args) and, if it fails, automatically tries
        fallback_fn(*args) instead.

        It returns (result, used_fallback) where used_fallback is True
        if the fallback was needed. This lets us log which path was taken.

        This is cleaner than writing:
            try:
                return _extract_pdf_with_pdfplumber(data)
            except:
                return _extract_pdf_with_pypdf2(data)
        Because with_fallback handles logging, categorization, and the
        used_fallback flag all in one place.

    📌 TEACHING NOTE — 'raise ... from e' (exception chaining):
        raise FileParsingError("...message...") from e

        This creates an exception CHAIN:
        - The new FileParsingError is what the caller sees
        - The original exception 'e' is attached as the cause
        - Python will show both when printing the traceback

        This preserves the root cause for debugging while exposing a
        cleaner, user-friendly error to the caller.
        Without 'from e', the original cause would be hidden.

    Args:
        file_data: Raw PDF bytes

    Returns:
        Extracted text string

    Raises:
        FileParsingError: If both pdfplumber and PyPDF2 fail
    """
    try:
        result, used_fallback = with_fallback(
            _extract_pdf_with_pdfplumber,
            _extract_pdf_with_pypdf2,
            file_data,
            error_category=ErrorCategory.TEXT_EXTRACTION,
            log_fallback=True
        )
        if used_fallback:
            log_info('PDF extraction succeeded using PyPDF2 fallback', context='file_parser')
        return result
    except Exception as e:
        log_error(e, context='extract_text_from_pdf', category=ErrorCategory.TEXT_EXTRACTION)
        raise FileParsingError(
            'Failed to extract text from PDF using both pdfplumber and PyPDF2. '
            'The PDF may be corrupted, password-protected, or contain only images. '
            'Please try converting the PDF to a different format or ensure it contains selectable text.'
        ) from e   # 'from e' chains the original exception for debugging


# ── DOCX Text Extraction ──────────────────────────────────────────────────────

def extract_text_from_docx(file_data: bytes) -> str:
    """
    Extract text from a Microsoft Word .docx file.

    📌 TEACHING NOTE — DOCX file structure:
        A .docx file is actually a ZIP archive containing XML files.
        The python-docx library (from docx import Document) handles
        unzipping and parsing the XML so we just call paragraph.text.

        doc.paragraphs  → all paragraph elements (most text)
        doc.tables      → tables in the document (skills often in tables!)

        Important: Resume content is sometimes in tables for layout purposes.
        If we only read paragraphs, we'd miss table content entirely.
        Both are read here to ensure complete extraction.

    📌 TEACHING NOTE — text_parts list + join:
        Instead of building a string with +=, we append to a list and
        join at the end. This is O(n) vs O(n²) for string concatenation.
        (See generator.py teaching note for full explanation)

    📌 TEACHING NOTE — raise ... / except FileParsingError: raise:
        The outer except catches all exceptions.
        But we re-raise FileParsingError without re-wrapping it
        (we want the original FileParsingError to propagate unchanged).

        Without this:
            except Exception as e:
                raise FileParsingError("...") from e
        Would wrap a FileParsingError inside another FileParsingError.

        Pattern: re-raise the specific exception you care about,
        catch-and-wrap everything else.

    Args:
        file_data: Raw DOCX bytes

    Returns:
        Extracted text as a string

    Raises:
        FileParsingError: If extraction fails
    """
    try:
        doc = Document(io.BytesIO(file_data))  # Parse the DOCX ZIP/XML structure
        text_parts = []

        # Read all paragraphs (the main body text)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():   # Skip empty paragraphs
                text_parts.append(paragraph.text)

        # Read all table cells (resume skills are often in tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        text = '\n'.join(text_parts)  # Efficient string building

        if not text.strip():
            log_warning('DOCX file contained no extractable text', context='extract_text_from_docx')
            raise FileParsingError(
                'No text could be extracted from the document. '
                'The document may be empty or corrupted. Please check the file and try again.'
            )

        log_info(f'Successfully extracted {len(text)} characters from DOCX', context='extract_text_from_docx')
        return text.strip()

    except FileParsingError:
        raise   # Re-raise as-is — don't wrap in another FileParsingError

    except Exception as e:
        log_error(e, context='extract_text_from_docx', category=ErrorCategory.TEXT_EXTRACTION)
        raise FileParsingError(
            'Failed to extract text from DOCX file. The document may be corrupted or in an unsupported format. '
            'Please try re-saving the document or converting it to PDF.'
        ) from e


# ── DOC Text Extraction ───────────────────────────────────────────────────────

def extract_text_from_doc(file_data: bytes) -> str:
    """
    Attempt to extract text from a legacy .doc file — intentionally unsupported.

    📌 TEACHING NOTE — Explicit NotImplemented vs silent failure:
        This function immediately raises an error instead of trying to parse.
        This is a deliberate design decision: .doc (binary format from Word 97-2003)
        is very complex to parse correctly in pure Python.

        Options considered:
        1. Try to parse and produce garbage text → misleads the user
        2. Use antiword/LibreOffice as system tool → adds complex dependencies
        3. Raise a clear error with actionable advice → chosen approach ✅

        Sometimes the best implementation is an honest "not supported" with
        clear instructions on what to do instead. This is better than a
        broken implementation that gives wrong results silently.

    📌 TEACHING NOTE — User-facing error messages should be actionable:
        Bad:  "doc format not supported"
        Good: "Legacy .doc format is not fully supported. Please convert
               your document to .docx or .pdf format and try again. You can
               convert using Microsoft Word, Google Docs, or online tools."

        The good version tells the user EXACTLY what to do next.
        Error messages are a UX concern, not just a technical one.
    """
    raise FileParsingError(
        'Legacy .doc format is not fully supported. '
        'Please convert your document to .docx or .pdf format and try again. '
        'You can convert using Microsoft Word, Google Docs, or online conversion tools.'
    )


# ── Dispatch Function ─────────────────────────────────────────────────────────

def extract_text(file_data: bytes, file_type: str) -> str:
    """
    Route file extraction to the correct function based on file type.

    📌 TEACHING NOTE — Dispatcher / Router Pattern:
        This function doesn't do any parsing itself — it just routes
        to the appropriate function based on file_type.

        This is the "Dispatcher Pattern":
            if file_type == 'pdf':  → extract_text_from_pdf()
            elif file_type == 'docx': → extract_text_from_docx()
            elif file_type == 'doc':  → extract_text_from_doc()

        Benefits:
        - Callers use ONE function regardless of file type
        - Adding a new format (e.g., 'txt') = add one elif here
        - Each format's logic stays in its own function (separation of concerns)

        Alternative: a dict mapping {'pdf': fn, 'docx': fn2, ...}
        then calling handler[file_type](file_data).
        Both approaches are valid — the if/elif is more readable for beginners.

    Args:
        file_data: Raw file bytes
        file_type: 'pdf', 'docx', or 'doc'

    Returns:
        Extracted text string

    Raises:
        FileValidationError: If file_type is not recognized
    """
    if file_type == 'pdf':
        return extract_text_from_pdf(file_data)
    elif file_type == 'docx':
        return extract_text_from_docx(file_data)
    elif file_type == 'doc':
        return extract_text_from_doc(file_data)
    else:
        raise FileValidationError(
            f'Invalid file type: {file_type}. Supported types are: PDF, DOC, DOCX.'
        )


# ── Main Public Entry Point ───────────────────────────────────────────────────

def parse_resume_file(file_data: bytes, filename: str) -> Tuple[str, dict]:
    """
    Top-level function — validates and extracts text from an uploaded resume.

    📌 TEACHING NOTE — Two separate try/except blocks:
        This function has TWO separate try/except blocks — one for validation,
        one for extraction. Why not one big try/except?

        Reason: different errors need different handling.
        - Validation failure → raise FileValidationError
        - Extraction failure → raise FileParsingError
        - Unexpected error in validation → wrap in FileValidationError
        - Unexpected error in extraction → wrap in FileParsingError

        If we used one block, we'd have to inspect the exception type to
        decide how to re-raise it — messier. Two blocks each focus on
        their own phase with clear responsibility.

    📌 TEACHING NOTE — 'except FileParsingError: raise' pattern again:
        Both try/except blocks re-raise their "expected" exception type
        unchanged, while wrapping any "unexpected" exceptions.

        This ensures the caller always gets a meaningful, typed exception:
        FileValidationError (user did something wrong with the file)
        FileParsingError (the file is valid but unreadable)
        ...never a bare AttributeError or KeyError from internal bugs.

    📌 TEACHING NOTE — Metadata dict:
        The function returns BOTH the text AND a metadata dict.
        The metadata captures facts about the parsing process:
        filename, file_type, file_size_bytes, text_length, success flag.

        This metadata can be stored in the database, shown in the UI,
        or used for debugging. It doesn't cost much to collect and
        can be very useful later.

    📌 TEACHING NOTE — log_info/log_warning/log_error calls:
        The function logs at different levels throughout:
        log_info()    → normal operation (started, succeeded)
        log_warning() → unexpected but non-fatal (validation failed)
        log_error()   → unexpected exception

        This logging enables debugging in production without the user
        having to reproduce the issue. Always log at entry and exit of
        important operations.

    Args:
        file_data: Raw bytes of the uploaded file
        filename: Original filename string

    Returns:
        Tuple of (extracted_text, metadata_dict)

    Raises:
        FileValidationError: If file is wrong type/size
        FileParsingError: If text extraction fails
    """
    log_info(f'Starting file parsing for: {filename}', context='parse_resume_file')

    # ── Phase 1: Validation ───────────────────────────────────────────────
    try:
        is_valid, error_msg, file_type = validate_file(file_data, filename)
        if not is_valid:
            log_warning(f'File validation failed: {error_msg}', context='parse_resume_file')
            raise FileValidationError(error_msg)

    except FileValidationError:
        raise   # Re-raise unchanged — this is the expected failure path

    except Exception as e:
        # Unexpected error during validation (shouldn't happen, but guard it)
        log_error(e, context='parse_resume_file_validation', category=ErrorCategory.FILE_UPLOAD)
        raise FileValidationError(
            'Could not validate the uploaded file. Please ensure it is a valid PDF, DOC, or DOCX file.'
        ) from e

    # ── Phase 2: Text extraction ──────────────────────────────────────────
    try:
        text = extract_text(file_data, file_type)
        log_info(
            f'Successfully extracted {len(text)} characters from {filename}',
            context='parse_resume_file'
        )

    except FileParsingError:
        raise   # Re-raise unchanged — this is the expected failure path

    except Exception as e:
        # Unexpected error during extraction (shouldn't happen, but guard it)
        log_error(e, context='parse_resume_file_extraction', category=ErrorCategory.TEXT_EXTRACTION)
        raise FileParsingError(
            'An unexpected error occurred while processing the file. '
            'Please try again or contact support if the problem persists.'
        ) from e

    # ── Build and return metadata ─────────────────────────────────────────
    metadata = {
        'filename':        filename,
        'file_type':       file_type,
        'file_size_bytes': len(file_data),
        'text_length':     len(text),
        'success':         True
    }
    return text, metadata